"""
generate_diagrams_tva_transfert.py
Generates 2 flow diagrams and inserts them into Manuel_Utilisation_SGCI_v4.docx:
  - diag_tva.png      → §15.3  after "La DGTCP vérifie, valide, puis apure."
  - diag_transfert.png → §15.4 after "Le Président valide ou rejette."
"""

import shutil
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import numpy as np

from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR  = Path(__file__).resolve().parent.parent
DIAG_DIR  = BASE_DIR / "scripts" / "diagrams"
DIAG_DIR.mkdir(exist_ok=True)

SRC_DOCX  = BASE_DIR / "Manuel_Utilisation_SGCI_v4.docx"
OUT_DOCX  = BASE_DIR / "Manuel_Utilisation_SGCI_v4_v2.docx"
BAK_DOCX  = BASE_DIR / "Manuel_Utilisation_SGCI_v4.bak3.docx"

# ── Palette ──────────────────────────────────────────────────────────────────
NAVY   = "#1F497D"
BLUE   = "#2E74B5"
TEAL   = "#1F7489"
ORANGE = "#C55A11"
PURPLE = "#7030A0"
RED    = "#C00000"
GREEN  = "#375623"
AMBER  = "#D4A017"
LGRAY  = "#F2F2F2"
WHITE  = "#FFFFFF"
DGRAY  = "#595959"


def _save(fig, name, dpi=180):
    path = DIAG_DIR / name
    fig.savefig(path, dpi=dpi, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    print(f"  saved {path.name}")
    return path


# ─────────────────────────────────────────────────────────────────────────────
# DIAGRAM TVA — Flux utilisation TVA intérieure
# ─────────────────────────────────────────────────────────────────────────────
def make_diag_tva():
    steps = [
        (GREEN,  "Entreprise",
         "Déclare une utilisation TVA\n(facture, montant, type)",
         "BROUILLON → DEMANDEE"),
        (ORANGE, "DGI",
         "Instruit le dossier\n(décision ou demande de compléments)",
         "EN_VERIFICATION"),
        (TEAL,   "DGTCP",
         "Vérifie et valide\n(mise à jour du solde TVA)",
         "VALIDEE"),
        (TEAL,   "DGTCP",
         "Apure la TVA intérieure\n(régularisation finale)",
         "APUREE"),
    ]

    # Side branch for INCOMPLETE / A_RECONTROLER
    branch_steps = [
        (AMBER, "Pièces\nmanquantes", "INCOMPLETE"),
        (BLUE,  "Compléments\nfournis",   "A_RECONTROLER"),
    ]

    fig_h = 1.0 + len(steps) * 1.15
    fig, ax = plt.subplots(figsize=(11, fig_h), facecolor=LGRAY)
    ax.set_xlim(0, 11)
    ax.set_ylim(0, fig_h)
    ax.axis("off")

    ax.text(5.5, fig_h - 0.35,
            "Flux : Utilisation du crédit — Circuit TVA intérieure",
            ha="center", va="center", fontsize=11, fontweight="bold", color=NAVY)

    line_x  = 2.0
    y_start = fig_h - 0.95

    # Vertical spine
    ax.plot([line_x, line_x], [0.35, y_start],
            color=ORANGE, lw=2.5, zorder=1, solid_capstyle="round")

    for i, (color, actor, action, status) in enumerate(steps):
        y = y_start - i * 1.15

        # Circle on spine
        circle = plt.Circle((line_x, y), 0.22, color=color, zorder=3,
                             linewidth=2, edgecolor=WHITE)
        ax.add_patch(circle)
        ax.text(line_x, y, str(i + 1),
                ha="center", va="center", fontsize=8, fontweight="bold",
                color=WHITE, zorder=4)

        # Connector to box
        ax.plot([line_x + 0.22, line_x + 0.55], [y, y],
                color=color, lw=1.5, zorder=2)

        # Content box
        box = FancyBboxPatch((line_x + 0.55, y - 0.38), 7.9, 0.76,
                             boxstyle="round,pad=0.06",
                             facecolor=WHITE, edgecolor=color,
                             linewidth=1.8, zorder=3)
        ax.add_patch(box)

        # Actor badge
        badge = FancyBboxPatch((line_x + 0.63, y - 0.31), 1.3, 0.62,
                               boxstyle="round,pad=0.04",
                               facecolor=color, edgecolor="none", zorder=4)
        ax.add_patch(badge)
        ax.text(line_x + 1.28, y, actor,
                ha="center", va="center", fontsize=7.5, fontweight="bold",
                color=WHITE, zorder=5)

        # Action text
        ax.text(line_x + 2.1, y + 0.07, action,
                ha="left", va="center", fontsize=8, color=DGRAY,
                linespacing=1.25, zorder=4)

        # Status chip (right)
        ax.text(line_x + 8.25, y, status,
                ha="right", va="center", fontsize=7, color=color,
                fontweight="bold", style="italic", zorder=4)

    # ── Side branch: INCOMPLETE / A_RECONTROLER (between step 2 and 3) ──
    branch_x  = 9.6
    branch_y1 = y_start - 1 * 1.15   # level of step 2 (DGI)
    branch_y2 = y_start - 2 * 1.15   # level of step 3 (DGTCP)

    # Dashed line down the side
    ax.plot([branch_x, branch_x], [branch_y1 - 0.1, branch_y2 + 0.42],
            color=AMBER, lw=1.4, linestyle="--", zorder=2)

    for bi, (bcolor, blabel, bstatus) in enumerate(branch_steps):
        by = branch_y1 - 0.1 - bi * 0.55
        bchip = FancyBboxPatch((branch_x - 0.55, by - 0.22), 1.1, 0.44,
                               boxstyle="round,pad=0.04",
                               facecolor=bcolor, edgecolor="none",
                               alpha=0.85, zorder=4)
        ax.add_patch(bchip)
        ax.text(branch_x, by - 0.01, f"{blabel}\n{bstatus}",
                ha="center", va="center", fontsize=6.2, color=WHITE,
                fontweight="bold", linespacing=1.2, zorder=5)

    ax.annotate("",
                xy=(line_x + 0.22, branch_y2 + 0.35),
                xytext=(branch_x - 0.55, branch_y2 + 0.35),
                arrowprops=dict(arrowstyle="-|>", color=BLUE, lw=1.2,
                                mutation_scale=12))
    ax.text((line_x + branch_x) / 2 + 0.6, branch_y2 + 0.52,
            "retour au contrôle", ha="center", va="center",
            fontsize=6.5, color=BLUE, style="italic")

    fig.tight_layout(pad=0.3)
    return _save(fig, "diag_tva.png")


# ─────────────────────────────────────────────────────────────────────────────
# DIAGRAM TRANSFERT — Flux transfert de crédit
# ─────────────────────────────────────────────────────────────────────────────
def make_diag_transfert():
    fig, ax = plt.subplots(figsize=(13, 5.8), facecolor=LGRAY)
    ax.set_xlim(0, 13); ax.set_ylim(0, 5.8)
    ax.axis("off")

    ax.text(6.5, 5.4,
            "Flux : Demande de transfert de crédit",
            ha="center", va="center", fontsize=11, fontweight="bold", color=NAVY)

    # ── Main flow (left to right) ─────────────────────────────────────────────
    main_stages = [
        (GREEN,  "Entreprise", "Crée la\ndemande",        "DEMANDE"),
        (GREEN,  "Entreprise", "Dépose les\npièces oblig.", "EN_COURS"),
        (TEAL,   "DGTCP",     "Instruit\nle dossier",     "EN_COURS"),
        (RED,    "Président", "Valide ou\nrejette",        "décision\nfinale"),
    ]

    box_w, box_h = 2.1, 1.7
    gap  = 0.45
    y_main = 2.6
    xs = [0.4 + i * (box_w + gap) for i in range(len(main_stages))]

    for i, (color, actor, action, status) in enumerate(main_stages):
        x = xs[i]
        box = FancyBboxPatch((x, y_main), box_w, box_h,
                             boxstyle="round,pad=0.09",
                             facecolor=color, edgecolor=WHITE,
                             linewidth=1.8, zorder=3)
        ax.add_patch(box)

        # Step badge
        badge = plt.Circle((x + box_w/2, y_main + box_h + 0.22), 0.21,
                            color=DGRAY, zorder=5)
        ax.add_patch(badge)
        ax.text(x + box_w/2, y_main + box_h + 0.22, str(i + 1),
                ha="center", va="center", fontsize=8, fontweight="bold",
                color=WHITE, zorder=6)

        # Actor
        ax.text(x + box_w/2, y_main + box_h * 0.75, actor,
                ha="center", va="center", fontsize=8.5, fontweight="bold",
                color=WHITE, zorder=4)

        # Action
        ax.text(x + box_w/2, y_main + box_h * 0.48, action,
                ha="center", va="center", fontsize=7.5, color=WHITE,
                linespacing=1.3, zorder=4)

        # Status chip
        chip = FancyBboxPatch((x + 0.18, y_main + 0.1), box_w - 0.36, 0.52,
                              boxstyle="round,pad=0.04",
                              facecolor=WHITE, edgecolor="none",
                              alpha=0.22, zorder=4)
        ax.add_patch(chip)
        ax.text(x + box_w/2, y_main + 0.36, status,
                ha="center", va="center", fontsize=6.5, fontweight="bold",
                color=WHITE, linespacing=1.2, zorder=5)

        # Arrow to next main stage
        if i < len(main_stages) - 1:
            ax.annotate("", xy=(xs[i+1], y_main + box_h/2),
                        xytext=(x + box_w, y_main + box_h/2),
                        arrowprops=dict(arrowstyle="-|>", color=NAVY,
                                        lw=2, mutation_scale=18), zorder=2)

    # ── Outcome branches (below main flow) ───────────────────────────────────
    outcomes = [
        (xs[3] + box_w/2, "#375623", "TRANSFERE",
         "Transfert exécuté\nsur le certificat", 1.5),
        (xs[3] + box_w/2, RED,       "REJETE",
         "Refus définitif", 0.75),
        (xs[0] + box_w/2, AMBER,     "ANNULEE",
         "Annulation par\nl'entreprise", 0.75),
    ]

    # TRANSFERE (down from Président)
    ax.annotate("", xy=(xs[3]+box_w/2, y_main - 0.04),
                xytext=(xs[3]+box_w/2, y_main - 1.4),
                arrowprops=dict(arrowstyle="-|>", color="#375623",
                                lw=1.6, mutation_scale=14))
    ok_box = FancyBboxPatch((xs[3]+box_w/2 - 1.15, 0.1), 2.3, 1.2,
                            boxstyle="round,pad=0.07",
                            facecolor="#375623", edgecolor=WHITE,
                            linewidth=1.5, zorder=3)
    ax.add_patch(ok_box)
    ax.text(xs[3]+box_w/2, 0.72, "TRANSFERE",
            ha="center", va="center", fontsize=9, fontweight="bold",
            color=WHITE, zorder=4)
    ax.text(xs[3]+box_w/2, 0.30, "Transfert exécuté\nsur le certificat",
            ha="center", va="center", fontsize=7, color=WHITE,
            linespacing=1.3, zorder=4)

    # REJETE (diagonal right)
    rej_x = 12.0
    rej_y_top = y_main + 0.4
    ax.annotate("", xy=(rej_x, rej_y_top - 0.6),
                xytext=(xs[3] + box_w, rej_y_top),
                arrowprops=dict(arrowstyle="-|>", color=RED,
                                lw=1.4, mutation_scale=13))
    rej_box = FancyBboxPatch((rej_x - 0.6, rej_y_top - 1.25), 1.2, 0.88,
                             boxstyle="round,pad=0.06",
                             facecolor=RED, edgecolor=WHITE,
                             linewidth=1.4, zorder=3)
    ax.add_patch(rej_box)
    ax.text(rej_x, rej_y_top - 0.72, "REJETE",
            ha="center", va="center", fontsize=8, fontweight="bold",
            color=WHITE, zorder=4)
    ax.text(rej_x, rej_y_top - 1.05, "Refus\ndéfinitif",
            ha="center", va="center", fontsize=6.5, color=WHITE,
            linespacing=1.2, zorder=4)

    # ANNULEE (left, from Entreprise)
    ann_x = 0.1
    ann_y = y_main - 0.3
    ax.annotate("", xy=(ann_x + 0.85, ann_y - 0.35),
                xytext=(xs[0], ann_y),
                arrowprops=dict(arrowstyle="-|>", color=AMBER,
                                lw=1.4, mutation_scale=13))
    ann_box = FancyBboxPatch((ann_x - 0.05, ann_y - 1.45), 1.5, 0.88,
                             boxstyle="round,pad=0.06",
                             facecolor=AMBER, edgecolor=WHITE,
                             linewidth=1.4, zorder=3)
    ax.add_patch(ann_box)
    ax.text(ann_x + 0.7, ann_y - 0.98, "ANNULEE",
            ha="center", va="center", fontsize=8, fontweight="bold",
            color=WHITE, zorder=4)
    ax.text(ann_x + 0.7, ann_y - 1.33, "Avant exécution",
            ha="center", va="center", fontsize=6.5, color=WHITE,
            linespacing=1.2, zorder=4)

    # ── Rejet temporaire arc (DGTCP → INCOMPLETE → A_RECONTROLER) ────────────
    rt_x_mid = xs[2] + box_w/2
    ax.annotate("", xy=(xs[2], y_main + box_h * 0.6),
                xytext=(xs[2] + box_w * 0.7, y_main + box_h + 0.65),
                arrowprops=dict(arrowstyle="-|>", color=PURPLE,
                                lw=1.3, connectionstyle="arc3,rad=0.4",
                                mutation_scale=12))
    ax.text(rt_x_mid + 0.3, y_main + box_h + 0.85,
            "Rejet temporaire (INCOMPLETE → compléments → A_RECONTROLER)",
            ha="center", va="center", fontsize=7, color=PURPLE,
            style="italic")

    fig.tight_layout(pad=0.3)
    return _save(fig, "diag_transfert.png")


# ─────────────────────────────────────────────────────────────────────────────
# INSERT INTO WORD
# ─────────────────────────────────────────────────────────────────────────────

def _insert_image_after_para(doc, para_index, img_path, caption, width_cm=14.0):
    body    = doc.element.body
    paras   = [el for el in body if el.tag != qn("w:sectPr")]
    if para_index >= len(paras):
        para_index = len(paras) - 1
    ref_el  = paras[para_index]

    img_para = doc.add_paragraph()
    img_para.alignment = 1
    img_para.paragraph_format.space_before = Pt(8)
    img_para.paragraph_format.space_after  = Pt(2)
    img_para.add_run().add_picture(str(img_path), width=Cm(width_cm))

    cap_para = doc.add_paragraph()
    cap_para.alignment = 1
    cap_para.paragraph_format.space_before = Pt(1)
    cap_para.paragraph_format.space_after  = Pt(10)
    cr = cap_para.add_run(caption); cr.italic = True; cr.font.size = Pt(9)

    body.remove(img_para._p); body.remove(cap_para._p)
    ref_el.addnext(cap_para._p)
    ref_el.addnext(img_para._p)


def insert_into_docx(diag_paths):
    if not SRC_DOCX.exists():
        print(f"Source not found: {SRC_DOCX}"); return

    shutil.copy2(SRC_DOCX, BAK_DOCX)
    doc  = Document(SRC_DOCX)
    body = doc.element.body
    paras = [el for el in body if el.tag != qn("w:sectPr")]

    para_texts = []
    for el in paras:
        para_texts.append(
            "".join(n.text or "" for n in el.iter() if n.tag == qn("w:t")).strip()
        )

    insertions = {
        "tva":      ("La DGTCP vérifie, valide, puis apure.",
                     diag_paths["tva"],
                     "Figure 8 \u2013 Flux : Utilisation TVA int\u00e9rieure"),
        "transfert":("Le Pr\u00e9sident valide ou rejette.",
                     diag_paths["transfert"],
                     "Figure 9 \u2013 Flux : Transfert de cr\u00e9dit"),
    }

    offset = 0
    for key, (search, img_path, caption) in insertions.items():
        found_idx = next(
            (i for i, t in enumerate(para_texts) if search in t), None
        )
        if found_idx is None:
            print(f"  [WARN] marker not found for {key}: '{search}'")
            continue
        actual = found_idx + offset
        _insert_image_after_para(doc, actual, img_path, caption)
        offset += 2
        preview = para_texts[found_idx][:70].encode("ascii", errors="replace").decode()
        print(f"  inserted {key} after para {found_idx} ('{preview}')")

    doc.save(OUT_DOCX)
    print(f"\nOK - Saved: {OUT_DOCX.name}")


# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Generating diagrams...")
    p_tva      = make_diag_tva()
    p_transfert = make_diag_transfert()
    print("\nInserting into Word document...")
    insert_into_docx({"tva": p_tva, "transfert": p_transfert})
    print("\nDone.")
