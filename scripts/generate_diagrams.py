"""
generate_diagrams.py
Generates 3 flow diagrams as PNG files and inserts them into the Word document.

Diagrams:
  - diag6_circuit_correction.png        → §4.4 after "circuit de visa"
  - diag_flux_correction_certificat.png → correction → mise en place → OUVERT
  - diag13_flux_convention.png          → §15.1 after the ASCII block
  - diag14_flux_douane.png              → §15.2 after the ASCII block
"""

import shutil
from pathlib import Path

import matplotlib
matplotlib.use("Agg")          # no display needed
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import matplotlib.patheffects as pe
import numpy as np

from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR  = Path(__file__).resolve().parent.parent
DIAG_DIR  = BASE_DIR / "scripts" / "diagrams"
DIAG_DIR.mkdir(exist_ok=True)

SRC_DOCX  = BASE_DIR / "Manuel_Utilisation_SGCI_v4.docx"
OUT_DOCX  = BASE_DIR / "Manuel_Utilisation_SGCI_v4_updated.docx"
BAK_DOCX  = BASE_DIR / "Manuel_Utilisation_SGCI_v4.bak2.docx"

# ── Palette ──────────────────────────────────────────────────────────────────
NAVY    = "#1F497D"
BLUE    = "#2E74B5"
TEAL    = "#1F7489"
ORANGE  = "#C55A11"
PURPLE  = "#7030A0"
RED     = "#C00000"
GREEN   = "#375623"
LGRAY   = "#F2F2F2"
WHITE   = "#FFFFFF"
DGRAY   = "#595959"


def _save(fig, name, dpi=180):
    path = DIAG_DIR / name
    fig.savefig(path, dpi=dpi, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    print(f"  saved {path.name}")
    return path


# ─────────────────────────────────────────────────────────────────────────────
# DIAGRAM 6 — Circuit de correction douanière
# ─────────────────────────────────────────────────────────────────────────────
def make_diag6():
    fig, ax = plt.subplots(figsize=(13, 4.2), facecolor=LGRAY)
    ax.set_xlim(0, 13); ax.set_ylim(0, 4.2)
    ax.axis("off")

    stages = [
        ("DGD", BLUE,   "Direction\nGénérale\ndes Douanes"),
        ("DGTCP", TEAL,  "Direction\ndu Trésor"),
        ("DGI",  ORANGE, "Direction\nGénérale\ndes Impôts"),
        ("DGB",  PURPLE, "Direction\ndu Budget"),
        ("Président", RED, "Président\nde la\nCommission"),
    ]
    box_w, box_h = 1.9, 1.55
    gap = 0.52
    xs = [0.55 + i*(box_w + gap) for i in range(len(stages))]
    y_box = 1.3

    # Title
    ax.text(6.5, 3.75, "Circuit de visa — Demande de correction douanière",
            ha="center", va="center", fontsize=11, fontweight="bold",
            color=NAVY)

    # Boxes
    for i, (short, color, long_name) in enumerate(stages):
        x = xs[i]
        box = FancyBboxPatch((x, y_box), box_w, box_h,
                             boxstyle="round,pad=0.08",
                             facecolor=color, edgecolor=WHITE, linewidth=1.5)
        ax.add_patch(box)
        ax.text(x + box_w/2, y_box + box_h*0.68, short,
                ha="center", va="center", fontsize=11, fontweight="bold",
                color=WHITE)
        ax.text(x + box_w/2, y_box + box_h*0.25, long_name,
                ha="center", va="center", fontsize=7, color=WHITE, linespacing=1.3)
        # Step number badge
        badge = plt.Circle((x + box_w/2, y_box + box_h + 0.22), 0.18,
                            color=DGRAY, zorder=5)
        ax.add_patch(badge)
        ax.text(x + box_w/2, y_box + box_h + 0.22, str(i+1),
                ha="center", va="center", fontsize=8, fontweight="bold",
                color=WHITE, zorder=6)

    # Forward arrows between boxes
    for i in range(len(stages)-1):
        x_start = xs[i] + box_w + 0.04
        x_end   = xs[i+1] - 0.04
        y_mid   = y_box + box_h/2
        ax.annotate("", xy=(x_end, y_mid), xytext=(x_start, y_mid),
                    arrowprops=dict(arrowstyle="-|>", color=NAVY,
                                   lw=1.8, mutation_scale=16))

    # AC box on the left (submitter)
    ac_x, ac_y = 0.05, 0.15
    ac_box = FancyBboxPatch((ac_x, ac_y), 1.3, 0.75,
                            boxstyle="round,pad=0.07",
                            facecolor=GREEN, edgecolor=WHITE, linewidth=1.2)
    ax.add_patch(ac_box)
    ax.text(ac_x + 0.65, ac_y + 0.38, "AC / UPM\n(soumission)",
            ha="center", va="center", fontsize=7.5, fontweight="bold",
            color=WHITE, linespacing=1.3)

    # Adoption box on the right
    ad_x = xs[-1] + box_w + 0.08
    ad_box = FancyBboxPatch((ad_x, ac_y), 1.3, 0.75,
                            boxstyle="round,pad=0.07",
                            facecolor=GREEN, edgecolor=WHITE, linewidth=1.2)
    ax.add_patch(ad_box)
    ax.text(ad_x + 0.65, ac_y + 0.38, "ADOPTEE\n→ NOTIFIEE",
            ha="center", va="center", fontsize=7.5, fontweight="bold",
            color=WHITE, linespacing=1.3)

    # Submission arrow (AC → DGD)
    ax.annotate("", xy=(xs[0], y_box + 0.35), xytext=(ac_x + 1.3, ac_y + 0.38),
                arrowprops=dict(arrowstyle="-|>", color=GREEN, lw=1.5,
                                connectionstyle="arc3,rad=-0.25", mutation_scale=14))

    # Adoption arrow (Président → ADOPTEE box)
    ax.annotate("", xy=(ad_x, ac_y + 0.38),
                xytext=(xs[-1] + box_w, y_box + 0.35),
                arrowprops=dict(arrowstyle="-|>", color=GREEN, lw=1.5,
                                connectionstyle="arc3,rad=-0.25", mutation_scale=14))

    # Rejet temporaire arc (bottom)
    ax.annotate("", xy=(xs[0], y_box),
                xytext=(xs[-1] + box_w*0.5, y_box),
                arrowprops=dict(arrowstyle="-|>", color="#C00000", lw=1.2,
                                connectionstyle="arc3,rad=0.55", mutation_scale=14))
    ax.text(6.5, 0.05, "Rejet temporaire / demande de compléments",
            ha="center", va="bottom", fontsize=7.5, color=RED,
            style="italic")

    fig.tight_layout(pad=0.3)
    return _save(fig, "diag6_circuit_correction.png")


def _draw_pipeline_row(ax, stages, y_box, box_w, box_h, gap, x0, step_offset=0):
    """Draw a horizontal pipeline row; status chips are placed below each card."""
    xs = []
    status_y = y_box - 0.62
    for i, (color, actor, action, status) in enumerate(stages):
        x = x0 + i * (box_w + gap)
        xs.append(x)
        cx = x + box_w / 2

        # Step badge — top-right inside the card (does not float above)
        badge = plt.Circle((x + box_w - 0.22, y_box + box_h - 0.22), 0.17,
                           color=DGRAY, zorder=6)
        ax.add_patch(badge)
        ax.text(x + box_w - 0.22, y_box + box_h - 0.22, str(step_offset + i + 1),
                ha="center", va="center", fontsize=7.5, fontweight="bold",
                color=WHITE, zorder=7)

        box = FancyBboxPatch((x, y_box), box_w, box_h,
                             boxstyle="round,pad=0.1",
                             facecolor=color, edgecolor=WHITE, linewidth=1.8,
                             zorder=3)
        ax.add_patch(box)

        ax.text(cx, y_box + box_h * 0.72, actor,
                ha="center", va="center", fontsize=8, fontweight="bold",
                color=WHITE, linespacing=1.2, zorder=4)
        ax.text(cx, y_box + box_h * 0.38, action,
                ha="center", va="center", fontsize=7.5, color=WHITE,
                linespacing=1.2, zorder=4)

        # Status chip below the card (always visible, not clipped)
        chip_w = box_w - 0.08
        chip = FancyBboxPatch((x + 0.04, status_y), chip_w, 0.52,
                              boxstyle="round,pad=0.04",
                              facecolor=WHITE, edgecolor=color,
                              linewidth=1.2, zorder=4)
        ax.add_patch(chip)
        ax.text(cx, status_y + 0.26, status,
                ha="center", va="center", fontsize=6.8, fontweight="bold",
                color=color, linespacing=1.15, zorder=5)

        if i < len(stages) - 1:
            y_mid = y_box + box_h / 2
            ax.annotate("", xy=(x + box_w + gap, y_mid),
                        xytext=(x + box_w, y_mid),
                        arrowprops=dict(arrowstyle="-|>", color=NAVY,
                                        lw=1.8, mutation_scale=16),
                        zorder=2)
    return xs


# ─────────────────────────────────────────────────────────────────────────────
# DIAGRAM — Flux correction → mise en place → certificat OUVERT
# ─────────────────────────────────────────────────────────────────────────────
def make_diag_flux_correction_certificat():
    fig, ax = plt.subplots(figsize=(16, 9.5), facecolor=LGRAY)
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 9.5)
    ax.axis("off")

    ax.text(8, 9.05,
            "Flux : Demande de correction → Mise en place → Certificat de crédit (OUVERT)",
            ha="center", va="center", fontsize=12, fontweight="bold", color=NAVY)

    box_w, box_h, gap = 2.15, 1.75, 0.45

    phase1 = [
        (GREEN,  "AC / UPM / UEP", "Crée et soumet\nla correction", "BROUILLON → RECUE"),
        (BLUE,   "DGD → DGTCP\n→ DGI → DGB", "Visas séquentiels\n(rejet temp. possible)", "EN_EVALUATION\n→ EN_VALIDATION"),
        (RED,    "Président", "Adopte la demande", "ADOPTEE"),
        (GREEN,  "Système", "Notification\naux parties", "NOTIFIEE"),
    ]
    phase2 = [
        (GREEN,  "AC / UPM / UEP", "Crée le certificat\n(lié à la correction)", "BROUILLON → ENVOYEE"),
        (TEAL,   "DGI / DGD / DGTCP", "Prise en charge\n+ visas (3 acteurs)", "EN_CONTROLE"),
        (RED,    "Président", "Valide et signe", "EN_VALIDATION_PRESIDENT\n→ VALIDE_PRESIDENT"),
        (NAVY,   "DGTCP", "Fixe montants\n+ ouvre le crédit", "EN_OUVERTURE_DGTCP\n→ OUVERT"),
    ]

    n = len(phase1)
    total_w = n * box_w + (n - 1) * gap
    x0 = (16 - total_w) / 2

    # Phase headers — above each row, never overlapping cards
    y1_header, y1_boxes = 6.55, 4.55
    y2_header, y2_boxes = 2.05, 1.05

    for y_hdr, label, color in [
        (y1_header, "1. Demande de correction douanière", ORANGE),
        (y2_header, "2. Demande de mise en place du certificat de crédit", TEAL),
    ]:
        hdr = FancyBboxPatch((x0 - 0.15, y_hdr - 0.18), total_w + 0.3, 0.36,
                             boxstyle="round,pad=0.04",
                             facecolor=color, edgecolor="none", alpha=0.15, zorder=1)
        ax.add_patch(hdr)
        ax.text(x0, y_hdr, label,
                ha="left", va="center", fontsize=9.5, fontweight="bold", color=color)

    xs1 = _draw_pipeline_row(ax, phase1, y1_boxes, box_w, box_h, gap, x0, step_offset=0)
    xs2 = _draw_pipeline_row(ax, phase2, y2_boxes, box_w, box_h, gap, x0, step_offset=4)

    # Bridge: NOTIFIEE (step 4) → step 5 — vertical corridor between rows
    x_last = xs1[-1] + box_w / 2
    x_first = xs2[0] + box_w / 2
    bridge_x = (x_last + x_first) / 2
    y_top = y1_boxes
    y_bot = y2_boxes + box_h
    ax.plot([bridge_x, bridge_x], [y_top, y_bot], color=PURPLE, lw=2, zorder=2)
    ax.annotate("", xy=(bridge_x, y_bot), xytext=(bridge_x, y_top),
                arrowprops=dict(arrowstyle="-|>", color=PURPLE, lw=2.2, mutation_scale=18),
                zorder=2)
    ax.text(bridge_x + 0.15, (y_top + y_bot) / 2,
            "Demande adoptée\n→ mise en place",
            ha="left", va="center", fontsize=8.5, color=PURPLE,
            fontweight="bold", linespacing=1.25,
            bbox=dict(boxstyle="round,pad=0.35", facecolor=WHITE,
                      edgecolor=PURPLE, alpha=0.95))

    # Rejet temporaire — arcs and labels below status chips, not under cards
    def _rejet_arc(xs, y_boxes, label):
        y_arc = y_boxes - 0.95
        ax.annotate("", xy=(xs[0] + 0.1, y_boxes),
                    xytext=(xs[-1] + box_w - 0.1, y_boxes),
                    arrowprops=dict(arrowstyle="-|>", color=RED, lw=1.2,
                                    connectionstyle="arc3,rad=-0.55", mutation_scale=12))
        ax.text(x0 + total_w / 2, y_arc, label,
                ha="center", va="center", fontsize=7.5, color=RED, style="italic",
                bbox=dict(boxstyle="round,pad=0.25", facecolor=WHITE,
                          edgecolor=RED, alpha=0.85, linewidth=0.8))

    _rejet_arc(xs1, y1_boxes,
               "Rejet temporaire / compléments : INCOMPLETE → A_RECONTROLER")
    _rejet_arc(xs2, y2_boxes,
               "Rejet temporaire certificat : INCOMPLETE → A_RECONTROLER")

    # Final outcome — right of last card, aligned with phase 2
    out_x = xs2[-1] + box_w + 0.25
    out_y = y2_boxes + box_h / 2
    out = FancyBboxPatch((out_x, out_y - 0.45), 1.65, 0.9,
                         boxstyle="round,pad=0.06",
                         facecolor=NAVY, edgecolor=WHITE, linewidth=1.5, zorder=3)
    ax.add_patch(out)
    ax.text(out_x + 0.825, out_y, "Certificat\nOUVERT",
            ha="center", va="center", fontsize=9, fontweight="bold",
            color=WHITE, linespacing=1.2)
    ax.annotate("", xy=(out_x, out_y), xytext=(xs2[-1] + box_w, out_y),
                arrowprops=dict(arrowstyle="-|>", color=NAVY, lw=1.8, mutation_scale=16))

    fig.tight_layout(pad=0.4)
    return _save(fig, "diag_flux_correction_certificat.png", dpi=200)


# ─────────────────────────────────────────────────────────────────────────────
# DIAGRAM 13 — Flux convention → certificat (pipeline)
# ─────────────────────────────────────────────────────────────────────────────
def make_diag13():
    fig, ax = plt.subplots(figsize=(14, 5.5), facecolor=LGRAY)
    ax.set_xlim(0, 14); ax.set_ylim(0, 5.5)
    ax.axis("off")

    ax.text(7, 5.1, "Flux : Convention → Mise en place du Certificat de crédit",
            ha="center", va="center", fontsize=11, fontweight="bold", color=NAVY)

    stages = [
        (GREEN,  "AC",          "Crée la\nconvention",         "EN_ATTENTE"),
        (PURPLE, "DGB/DGI\nPrésid.", "Valident la\nconvention",    "VALIDE"),
        (BLUE,   "AC / UPM\n/ UEP", "Créent le\ncertificat",    "BROUILLON\n→ ENVOYEE"),
        (TEAL,   "DGI / DGD\n/ DGTCP", "Contrôlent\nles pièces", "EN_CONTROLE"),
        (RED,    "Président",  "Valide et\nsigne",              "VALIDE_\nPRESIDENT"),
        (NAVY,   "DGTCP",     "Ouvre le\ncrédit",              "OUVERT"),
    ]

    box_w, box_h = 1.85, 2.0
    gap = 0.38
    total = len(stages) * box_w + (len(stages)-1) * gap
    x0 = (14 - total) / 2
    y_box = 0.8

    for i, (color, actor, action, status) in enumerate(stages):
        x = x0 + i * (box_w + gap)
        # Main box
        box = FancyBboxPatch((x, y_box), box_w, box_h,
                             boxstyle="round,pad=0.1",
                             facecolor=color, edgecolor=WHITE, linewidth=1.8,
                             zorder=3)
        ax.add_patch(box)
        # Step badge
        badge = plt.Circle((x + box_w/2, y_box + box_h + 0.22), 0.2,
                            color=DGRAY, zorder=5)
        ax.add_patch(badge)
        ax.text(x + box_w/2, y_box + box_h + 0.22, str(i+1),
                ha="center", va="center", fontsize=8.5, fontweight="bold",
                color=WHITE, zorder=6)
        # Actor label (top of box)
        ax.text(x + box_w/2, y_box + box_h*0.78, actor,
                ha="center", va="center", fontsize=8, fontweight="bold",
                color=WHITE, linespacing=1.3, zorder=4)
        # Action text (middle)
        ax.text(x + box_w/2, y_box + box_h*0.52, action,
                ha="center", va="center", fontsize=7.5, color=WHITE,
                linespacing=1.3, zorder=4)
        # Status chip at the bottom
        chip = FancyBboxPatch((x + 0.15, y_box + 0.1), box_w - 0.3, 0.52,
                              boxstyle="round,pad=0.05",
                              facecolor=WHITE, edgecolor="none",
                              alpha=0.25, zorder=4)
        ax.add_patch(chip)
        ax.text(x + box_w/2, y_box + 0.36, status,
                ha="center", va="center", fontsize=6.5, fontweight="bold",
                color=WHITE, linespacing=1.2, zorder=5)

        # Arrow to next
        if i < len(stages) - 1:
            ax.annotate("", xy=(x + box_w + gap, y_box + box_h/2),
                        xytext=(x + box_w, y_box + box_h/2),
                        arrowprops=dict(arrowstyle="-|>", color=NAVY,
                                        lw=2, mutation_scale=18),
                        zorder=2)

    fig.tight_layout(pad=0.3)
    return _save(fig, "diag13_flux_convention.png")


# ─────────────────────────────────────────────────────────────────────────────
# DIAGRAM 14 — Flux utilisation Douane (vertical timeline)
# ─────────────────────────────────────────────────────────────────────────────
def make_diag14():
    steps = [
        (GREEN,  "Entreprise",  "Soumet la demande",              "BROUILLON → DEMANDEE"),
        (BLUE,   "DGD",         "Contrôle et vise le bulletin\nde liquidation (AU_CI / À PAYER)", "EN_CONTROLE_DGD → VISE"),
        (GREEN,  "Entreprise",  "Saisit le chèque certifié\n(banque, N°, montant)",  "CHEQUE_SAISI"),
        (TEAL,   "DGTCP",       "Valide et envoie au Trésor",     "ENVOYEE_AU_TRESOR"),
        (TEAL,   "DGTCP",       "Saisit les quittances Trésor",   "QUITTANCES_ENREGISTREES"),
        (TEAL,   "DGTCP",       "Liquide — débit du solde\ncertificat d'utilisation généré", "LIQUIDEE"),
        (GREEN,  "Entreprise",  "Accuse réception",               "CLOTUREE"),
    ]

    fig_h = 1.1 + len(steps) * 1.05
    fig, ax = plt.subplots(figsize=(11, fig_h), facecolor=LGRAY)
    ax.set_xlim(0, 11)
    ax.set_ylim(0, fig_h)
    ax.axis("off")

    ax.text(5.5, fig_h - 0.35, "Flux : Utilisation du crédit — Circuit Douane",
            ha="center", va="center", fontsize=11, fontweight="bold", color=NAVY)

    line_x = 2.2
    y_start = fig_h - 1.0

    # Vertical spine
    ax.plot([line_x, line_x], [0.3, y_start],
            color=BLUE, lw=2.5, zorder=1, solid_capstyle="round")

    for i, (color, actor, action, status) in enumerate(steps):
        y = y_start - i * 1.05

        # Circle on spine
        circle = plt.Circle((line_x, y), 0.23, color=color, zorder=3, linewidth=2,
                             edgecolor=WHITE)
        ax.add_patch(circle)
        ax.text(line_x, y, str(i+1), ha="center", va="center",
                fontsize=8, fontweight="bold", color=WHITE, zorder=4)

        # Connector line to box
        ax.plot([line_x + 0.23, line_x + 0.6], [y, y],
                color=color, lw=1.5, zorder=2)

        # Content box
        box = FancyBboxPatch((line_x + 0.6, y - 0.36), 7.8, 0.72,
                             boxstyle="round,pad=0.06",
                             facecolor=WHITE, edgecolor=color,
                             linewidth=1.8, zorder=3)
        ax.add_patch(box)

        # Actor badge
        badge = FancyBboxPatch((line_x + 0.68, y - 0.3), 1.35, 0.6,
                               boxstyle="round,pad=0.04",
                               facecolor=color, edgecolor="none", zorder=4)
        ax.add_patch(badge)
        ax.text(line_x + 1.35, y, actor,
                ha="center", va="center", fontsize=7.5, fontweight="bold",
                color=WHITE, zorder=5)

        # Action text
        ax.text(line_x + 2.2, y + 0.09, action,
                ha="left", va="center", fontsize=8, color=DGRAY,
                linespacing=1.25, zorder=4)

        # Status chip
        ax.text(line_x + 8.1, y, status,
                ha="right", va="center", fontsize=7, color=color,
                fontweight="bold", style="italic", zorder=4)

    fig.tight_layout(pad=0.3)
    return _save(fig, "diag14_flux_douane.png")


# ─────────────────────────────────────────────────────────────────────────────
# INSERT DIAGRAMS INTO THE WORD DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────

INSERTION_MARKERS = {
    # (search_text, after=True, image_path_key, caption)
    "diag6":  ("DGD \u2192 DGTCP \u2192 DGI \u2192 DGB \u2192 Pr\u00e9sident",
               "diag6_circuit_correction.png",
               "Figure 5 \u2013 Circuit de visa d\u2019une demande de correction"),
    "diag13": ("DGTCP ouvre le cr\u00e9dit (EN_OUVERTURE_DGTCP \u2192 OUVERT)",
               "diag13_flux_convention.png",
               "Figure 6 \u2013 Flux : Convention \u2192 Mise en place du certificat"),
    "diag14": ("Entreprise accuse r\u00e9ception \u2192 CLOTUREE",
               "diag14_flux_douane.png",
               "Figure 7 \u2013 Flux : Utilisation du cr\u00e9dit \u2014 Circuit Douane"),
}


def _add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = 1   # center
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def _insert_image_after_para(doc, para_index, img_path, caption, width_cm=14.5):
    """Insert an image + caption after the paragraph at para_index."""
    from docx.oxml import OxmlElement as OE
    from copy import deepcopy

    # We build two new paragraphs (image + caption) and splice them in
    # by manipulating the body XML directly.
    body = doc.element.body

    # Find the paragraph element by counting non-sectPr block-level elements
    paras = [el for el in body if el.tag != qn("w:sectPr")]
    if para_index >= len(paras):
        para_index = len(paras) - 1
    ref_el = paras[para_index]

    # 1. Image paragraph
    img_para = doc.add_paragraph()
    img_para.alignment = 1
    img_para.paragraph_format.space_before = Pt(6)
    img_para.paragraph_format.space_after  = Pt(2)
    run = img_para.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))

    # 2. Caption paragraph
    cap_para = doc.add_paragraph()
    cap_para.alignment = 1
    cap_para.paragraph_format.space_before = Pt(1)
    cap_para.paragraph_format.space_after  = Pt(8)
    cr = cap_para.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(9)

    # Move both new elements to after ref_el
    img_el = img_para._p
    cap_el = cap_para._p
    body.remove(img_el)
    body.remove(cap_el)
    ref_el.addnext(cap_el)
    ref_el.addnext(img_el)


def insert_into_docx(diag_paths):
    if not SRC_DOCX.exists():
        print(f"Source not found: {SRC_DOCX}")
        return

    shutil.copy2(SRC_DOCX, BAK_DOCX)
    doc = Document(SRC_DOCX)
    body = doc.element.body
    paras = [el for el in body if el.tag != qn("w:sectPr")]

    # Build a text index for searching
    para_texts = []
    for el in paras:
        text = "".join(
            n.text or "" for n in el.iter() if n.tag == qn("w:t")
        )
        para_texts.append(text.strip())

    insertions = [
        # (search_text, img_key, caption, search_is_substring)
        ("DGD",        "diag6",  True),
        ("DGTCP ouvre le cr",    "diag13", True),
        ("Entreprise accuse r",  "diag14", True),
    ]

    # Map search hints to full marker info
    markers = {
        "diag6":  ("Les 4 directions apposent leurs visas",
                   diag_paths["diag6"],
                   "Figure 5 \u2013 Circuit de visa d\u2019une demande de correction"),
        "diag13": ("La DGTCP ouvre le certificat",
                   diag_paths["diag13"],
                   "Figure 6 \u2013 Flux : Convention \u2192 Certificat de cr\u00e9dit"),
        "diag14": ("La DGTCP liquide",
                   diag_paths["diag14"],
                   "Figure 7 \u2013 Flux : Utilisation Douane"),
    }

    # More precise searches
    precise = {
        "diag6":  "Les 4 directions apposent leurs visas",
        "diag13": "La DGTCP ouvre le certificat",
        "diag14": "La DGTCP liquide",
    }

    offset = 0   # each insertion shifts indices
    for key, (hint, img_path, caption) in markers.items():
        search = precise[key]
        found_idx = None
        for idx, txt in enumerate(para_texts):
            if search in txt:
                found_idx = idx
                break
        if found_idx is None:
            print(f"  [WARN] marker not found for {key}: '{search}'")
            continue
        actual_idx = found_idx + offset
        _insert_image_after_para(doc, actual_idx, img_path, caption)
        offset += 2   # each insertion adds 2 paragraphs
        preview = para_texts[found_idx][:60].encode("ascii", errors="replace").decode()
        print(f"  inserted {key} after para {found_idx} ('{preview}...')")

    doc.save(OUT_DOCX)
    print(f"\nOK - Saved: {OUT_DOCX.name}")


# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys

    only = sys.argv[1] if len(sys.argv) > 1 else None

    print("Generating diagrams...")
    generated = {}
    if only in (None, "correction-certificat", "all"):
        generated["correction_certificat"] = make_diag_flux_correction_certificat()
    if only in (None, "all", "6"):
        generated["diag6"] = make_diag6()
    if only in (None, "all", "13"):
        generated["diag13"] = make_diag13()
    if only in (None, "all", "14"):
        generated["diag14"] = make_diag14()

    if only is None and SRC_DOCX.exists():
        print("\nInserting into Word document...")
        insert_into_docx({
            "diag6": generated.get("diag6"),
            "diag13": generated.get("diag13"),
            "diag14": generated.get("diag14"),
        })
    elif only is None and not SRC_DOCX.exists():
        print(f"\nSkip Word insert — source not found: {SRC_DOCX.name}")

    print("\nDone.")
